#!/usr/bin/env python3
"""
生成符合中国版权保护中心计算机软件著作权申报标准的源代码文档
软件名称：青跃智衡 — AI姿态与动作筛查系统 V2.0
"""

import os
import sys

# ============================================================
# 配置
# ============================================================
SOFTWARE_NAME = "青跃智衡 — AI姿态与动作筛查系统 V2.0"
PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))

# 需要包含的源代码文件（相对 PROJECT_ROOT）
SOURCE_FILES = [
    # ---- 后端核心 ----
    "backend/app/__init__.py",
    "backend/app/main.py",
    "backend/app/core/__init__.py",
    "backend/app/core/config.py",
    "backend/app/core/errors.py",
    "backend/app/core/logging.py",
    "backend/app/shared/__init__.py",
    "backend/app/shared/db.py",
    # ---- 后端 API 层 ----
    "backend/app/api/__init__.py",
    "backend/app/api/deps.py",
    "backend/app/api/router.py",
    "backend/app/api/routes/__init__.py",
    "backend/app/api/routes/health.py",
    "backend/app/api/routes/assessments.py",
    "backend/app/api/routes/sessions.py",
    "backend/app/api/routes/reports.py",
    "backend/app/api/routes/camera.py",
    "backend/app/api/routes/screening.py",
    # ---- 后端 squat 功能模块 ----
    "backend/app/features/__init__.py",
    "backend/app/features/squat/__init__.py",
    "backend/app/features/squat/schemas.py",
    "backend/app/features/squat/service.py",
    "backend/app/features/squat/repository.py",
    "backend/app/features/squat/pose_contract.py",
    "backend/app/features/squat/live_analysis.py",
    "backend/app/features/squat/visual_detection.py",
    "backend/app/features/squat/visual_scoring.py",
    # ---- 后端 screening 功能模块 ----
    "backend/app/features/screening/__init__.py",
    "backend/app/features/screening/schemas.py",
    "backend/app/features/screening/service.py",
    "backend/app/features/screening/repository.py",
    # ---- 后端测试 ----
    "backend/tests/test_health.py",
    "backend/tests/test_squat_api.py",
    "backend/tests/test_camera.py",
    # ---- 前端入口与配置 ----
    "frontend/src/main.tsx",
    "frontend/src/App.tsx",
    "frontend/src/vite-env.d.ts",
    # ---- 前端共享模块 ----
    "frontend/src/shared/api/client.ts",
    "frontend/src/shared/types/api.ts",
    "frontend/src/shared/config/env.ts",
    "frontend/src/shared/components/Icon.tsx",
    "frontend/src/shared/components/LanguageSwitcher.tsx",
    "frontend/src/shared/components/ui/ProgressBar.tsx",
    "frontend/src/shared/components/ui/StatCard.tsx",
    "frontend/src/shared/components/ui/Button.tsx",
    "frontend/src/shared/components/ui/ActivityCard.tsx",
    "frontend/src/shared/components/ui/EmptyState.tsx",
    "frontend/src/shared/components/ui/DataBadge.tsx",
    "frontend/src/shared/components/ui/InsightCard.tsx",
    "frontend/src/shared/components/ui/SearchInput.tsx",
    "frontend/src/shared/components/ui/SurfaceCard.tsx",
    "frontend/src/shared/components/ui/RecordItem.tsx",
    "frontend/src/shared/components/ui/TopAppBar.tsx",
    "frontend/src/shared/components/ui/BottomNavBar.tsx",
    "frontend/src/shared/components/ui/index.ts",
    "frontend/src/shared/layout/AppShell.tsx",
    "frontend/src/shared/i18n/index.ts",
    # ---- 前端 squat 功能模块 ----
    "frontend/src/features/squat/components/ScoreCard.tsx",
    "frontend/src/features/squat/components/CameraFeed.tsx",
    "frontend/src/features/squat/pages/SquatSessionPage.tsx",
    "frontend/src/features/squat/components/CameraFeed.test.tsx",
    "frontend/src/features/squat/pages/SquatSessionPage.test.tsx",
    # ---- 前端 home 功能模块 ----
    "frontend/src/features/home/hooks/useHomeDashboard.ts",
    "frontend/src/features/home/pages/HomePage.tsx",
    # ---- 前端 history 功能模块 ----
    "frontend/src/features/history/hooks/useSessionSummaries.ts",
    "frontend/src/features/history/hooks/useSessionReportPreview.ts",
    "frontend/src/features/history/components/HistoryToolbar.tsx",
    "frontend/src/features/history/components/HistoryList.tsx",
    "frontend/src/features/history/components/HistoryEmptyState.tsx",
    "frontend/src/features/history/pages/HistoryPage.tsx",
    "frontend/src/features/history/pages/HistoryDetailPage.tsx",
    "frontend/src/features/history/lib/sessionUtils.ts",
    # ---- 前端 settings 功能模块 ----
    "frontend/src/features/settings/hooks/useSettingsOverview.ts",
    "frontend/src/features/settings/pages/SettingsPage.tsx",
    # ---- 前端测试 ----
    "frontend/src/test/setup.ts",
]


def collect_code_lines():
    """收集所有源代码行，返回列表[(file_path, code_lines)]"""
    all_files = []
    for rel_path in SOURCE_FILES:
        abs_path = os.path.join(PROJECT_ROOT, rel_path)
        if not os.path.exists(abs_path):
            print(f"  [跳过] 文件不存在: {rel_path}")
            continue
        with open(abs_path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
        lines = raw.split("\n")
        all_files.append((rel_path, lines))
        print(f"  [已读] {rel_path} ({len(lines)} 行)")
    return all_files


def generate_docx(all_files, output_path):
    """生成符合软著标准的 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("\n错误: 需要 python-docx 库。请运行: pip install python-docx")
        sys.exit(1)

    doc = Document()

    # ---- 页面设置：A4 ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.2)

    # ---- 设置默认字体为等宽字体 ----
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Courier New"
    font.size = Pt(7.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 段落格式：紧凑行距
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(10.5)

    # ---- 设置页眉 ----
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(4)

    # 页眉左侧：软件全称+版本号
    run_left = header_para.add_run(SOFTWARE_NAME)
    run_left.font.size = Pt(8)
    run_left.font.name = "宋体"
    run_left.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 页眉右侧：页码（使用 Word 域代码）
    # 右侧对齐制表位（页面宽度21cm - 右边距1.2cm = 19.8cm，减去左边距1.8cm = 18cm处）
    header_para.paragraph_format.tab_stops.add_tab_stop(Cm(18.0), WD_TAB_ALIGNMENT.RIGHT)

    # 使用 PAGE 域插入自动页码
    run_tab = header_para.add_run("\t")
    run_page = header_para.add_run()
    # 插入 PAGE 域
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run_page._r.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run_page._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_page._r.append(fldChar2)

    run_page.font.size = Pt(8)
    run_page.font.name = "宋体"
    run_page.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ---- 写入源代码内容 ----
    # 将所有代码行展开为统一列表
    all_lines = []
    for rel_path, code_lines in all_files:
        # 文件分隔标记
        all_lines.append(f"// ====== {rel_path} ======")
        all_lines.extend(code_lines)

    # 估算每页行数（含页眉空间） → A4 可用高度约 26cm，行距 10.5pt ≈ 0.37cm → 约 70 行/页
    LINES_PER_PAGE = 65

    # 分页写入
    total_lines = len(all_lines)
    page_num = 0
    idx = 0

    while idx < total_lines:
        page_num += 1
        chunk = all_lines[idx : idx + LINES_PER_PAGE]
        idx += LINES_PER_PAGE

        for line in chunk:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = Pt(10.5)
            # 处理空行：至少保留一个空格防止段落消失
            run = para.add_run(line if line.strip() else " ")
            run.font.name = "Courier New"
            run.font.size = Pt(7.5)
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ---- 保存 ----
    doc.save(output_path)
    print(f"\n✅ 文档已生成: {output_path}")
    print(f"   总页数: {page_num}")
    print(f"   总代码行数: {total_lines}")
    print(f"   包含文件数: {len(all_files)}")


def main():
    print("=" * 60)
    print(f"  软件著作权源代码文档生成器")
    print(f"  软件名称: {SOFTWARE_NAME}")
    print("=" * 60)

    print("\n[1/2] 收集源代码文件...")
    all_files = collect_code_lines()

    if not all_files:
        print("错误: 未找到任何源代码文件！")
        sys.exit(1)

    output_path = os.path.join(PROJECT_ROOT, f"{SOFTWARE_NAME}_源代码文档.docx")
    print(f"\n[2/2] 生成 Word 文档 → {output_path}")
    generate_docx(all_files, output_path)


if __name__ == "__main__":
    main()
